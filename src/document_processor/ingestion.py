import os
import io
import PyPDF2
import docx
from typing import Union, BinaryIO, Any, Optional, Dict, Tuple
import streamlit as st

from src.utils.logger import get_logger
from src.utils.azure_storage import AzureStorageManager

logger = get_logger(__name__)

class DocumentLoader:
    """
    Load and extract text from various document formats.
    Handles both file paths, Streamlit UploadedFile objects, and Azure Storage documents.
    """
    
    def __init__(self, use_azure: bool = False):
        """
        Initialize the document loader
        
        Args:
            use_azure: Whether to use Azure Storage for document operations
        """
        self.use_azure = use_azure
        self.azure_storage = AzureStorageManager() if use_azure else None
    
    def load_document(self, file_obj: Union[str, io.BytesIO, Any], doc_type: str = None, doc_id: str = None) -> Union[str, Tuple[str, Dict]]:
        """
        Extract text from a document file
        
        Args:
            file_obj: Can be a file path, BytesIO object, or Streamlit UploadedFile
            doc_type: Document type ('resume' or 'job') - required for Azure Storage
            doc_id: Document ID - required for Azure Storage
            
        Returns:
            Extracted text or tuple of (extracted text, metadata) if from Azure
        """
        # Handle Azure Storage documents if configured
        if self.use_azure and isinstance(file_obj, str) and doc_type and doc_id:
            return self._load_from_azure(doc_id, doc_type)
        
        try:
            # Handle Streamlit UploadedFile
            if hasattr(file_obj, 'name') and hasattr(file_obj, 'read'):
                filename = file_obj.name
                file_content = file_obj.read()
                
                # Create BytesIO object
                file_stream = io.BytesIO(file_content)
                
                # Reset file obj position for potential reuse
                if hasattr(file_obj, 'seek'):
                    file_obj.seek(0)
                
            # Handle file paths
            elif isinstance(file_obj, str):
                filename = os.path.basename(file_obj)
                with open(file_obj, 'rb') as f:
                    file_content = f.read()
                file_stream = io.BytesIO(file_content)
            
            # Handle BytesIO objects
            elif isinstance(file_obj, io.BytesIO):
                filename = "unknown"
                file_stream = file_obj
            
            else:
                raise TypeError(f"Unsupported file object type: {type(file_obj)}")
            
            # Get file extension
            _, ext = os.path.splitext(filename.lower())
            
            # Extract text based on file type
            if ext == '.pdf':
                return self._extract_from_pdf(file_stream)
            elif ext == '.docx':
                return self._extract_from_docx(file_stream)
            elif ext == '.txt' or ext == '.md':
                # Reset stream position
                file_stream.seek(0)
                return file_stream.read().decode('utf-8', errors='replace')
            else:
                raise ValueError(f"Unsupported file format: {ext}")
                
        except Exception as e:
            logger.error(f"Error loading document: {str(e)}")
            raise
    
    def _load_from_azure(self, doc_id: str, doc_type: str) -> Tuple[str, Dict]:
        """
        Load a document from Azure Storage
        
        Args:
            doc_id: Document ID
            doc_type: Type of document ('resume' or 'job')
            
        Returns:
            Tuple of (extracted text, metadata)
        """
        if not self.azure_storage:
            raise RuntimeError("Azure Storage not initialized")
        
        try:
            # Download document from Azure Storage
            file_content, metadata = self.azure_storage.download_document(doc_id, doc_type)
            
            # Create BytesIO object
            file_stream = io.BytesIO(file_content)
            
            # Extract text based on file type
            ext = metadata.get('file_extension', '').lower()
            
            if ext == '.pdf':
                text = self._extract_from_pdf(file_stream)
            elif ext == '.docx':
                text = self._extract_from_docx(file_stream)
            elif ext == '.txt' or ext == '.md':
                text = file_content.decode('utf-8', errors='replace')
            else:
                raise ValueError(f"Unsupported file format: {ext}")
                
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error loading document from Azure Storage: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_stream: BinaryIO) -> str:
        """
        Extract text from PDF
        
        Args:
            file_stream: File stream
            
        Returns:
            Extracted text
        """
        try:
            pdf_reader = PyPDF2.PdfReader(file_stream)
            text = ""
            
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"
                
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def _extract_from_docx(self, file_stream: BinaryIO) -> str:
        """
        Extract text from DOCX
        
        Args:
            file_stream: File stream
            
        Returns:
            Extracted text
        """
        try:
            doc = docx.Document(file_stream)
            text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                text += "\n"
                
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise